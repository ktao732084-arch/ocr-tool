import streamlit as st
import fitz  # PyMuPDF
from rapidocr_onnxruntime import RapidOCR
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn # 【核心修复】用于处理中文字体映射
import os
import tempfile
import time
import gc 

# --- 1. 页面配置 ---
st.set_page_config(page_title="终极OCR神器", page_icon="✨", layout="wide")

# 缓存模型，避免重复加载
@st.cache_resource
def load_model():
    return RapidOCR()

ocr_engine = load_model()

# --- 2. 侧边栏设置 ---
with st.sidebar:
    st.header("⚙️ 设置")
    st.markdown("针对中文乱码彻底修复")
    mode = st.radio("转换模式", ["🚀 极速预览 (推荐)", "🐢 高精模式"])
    
    if "极速" in mode:
        zoom_level = 1.5 # 平衡速度与清晰度
    else:
        zoom_level = 2.5

def process_pdf(pdf_path, docx_path, start_page, end_page, zoom):
    doc = fitz.open(pdf_path)
    word_doc = Document()
    
    # 设置全文档的基础样式，预防万一
    style = word_doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    total_pages = len(doc)
    start_idx = max(0, start_page - 1)
    end_idx = min(total_pages, end_page)
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("🖼️ 扫描页面")
        image_placeholder = st.empty()
    with col2:
        st.subheader("📝 识别结果")
        text_placeholder = st.empty()
        
    real_time_text = ""
    start_time = time.time()

    for i in range(start_idx, end_idx):
        page = doc[i]
        status_text.markdown(f"**正在处理第 {i + 1} 页...**")
        
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        img_bytes = pix.tobytes("png")
        
        # 【修复黄字警告】使用 use_container_width
        image_placeholder.image(img_bytes, caption=f"Page {i+1}", use_container_width=True)
        
        result, _ = ocr_engine(img_bytes)
        
        page_text = ""
        if result:
            for line in result:
                text = line[1].strip()
                if text:
                    # --- 【核心修复逻辑开始】 ---
                    # 不直接 add_paragraph(text)，而是分步处理字体
                    p = word_doc.add_paragraph()
                    run = p.add_run(text)
                    
                    # 1. 设置字体名称
                    run.font.name = 'Microsoft YaHei' # 使用微软雅黑
                    run.font.size = Pt(11)
                    
                    # 2. 【关键】显式设置东亚字体 XML 属性
                    # 这行代码告诉 Word："这是中文，请用微软雅黑显示，不要用 Arial 显示方框"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    # --- 【核心修复逻辑结束】 ---
                    
                    page_text += text + "\n"
            word_doc.add_page_break()
        
        real_time_text = f"--- 第 {i+1} 页 ---\n{page_text}\n" + real_time_text[:500]
        text_placeholder.text_area("识别预览", real_time_text, height=300)

        progress_bar.progress((i - start_idx + 1) / (end_idx - start_idx))
        
        del pix, img_bytes
        gc.collect()

    word_doc.save(docx_path)
    return time.time() - start_time, end_idx - start_idx

def main():
    st.title("✨ 终极版 PDF 转 Word")
    st.caption("✅ 已修复中文乱码方框问题 | ✅ 已移除黄字警告")

    uploaded_file = st.file_uploader("上传 PDF 文件", type="pdf")

    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
            tmp_pdf.write(uploaded_file.getvalue())
            tmp_pdf_path = tmp_pdf.name
        
        with fitz.open(tmp_pdf_path) as doc:
            total_pages = len(doc)

        st.success(f"文件加载成功！共 {total_pages} 页")

        c1, c2 = st.columns(2)
        with c1:
            start_p = st.number_input("开始页码", min_value=1, value=1)
        with c2:
            end_p = st.number_input("结束页码", min_value=1, value=min(5, total_pages), max_value=total_pages)
            
        if st.button("🚀 开始无乱码转换"):
            tmp_docx_path = tmp_pdf_path.replace(".pdf", ".docx")
            try:
                duration, pages = process_pdf(tmp_pdf_path, tmp_docx_path, start_p, end_p, zoom_level)
                st.balloons()
                st.success(f"🎉 成功！耗时 {duration:.2f} 秒")
                
                with open(tmp_docx_path, "rb") as file:
                    st.download_button(
                        label="📥 下载 Word (已修复字体)", 
                        data=file, 
                        file_name=f"Fixed_{uploaded_file.name.split('.')[0]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"出错: {str(e)}")
            finally:
                if os.path.exists(tmp_pdf_path):
                    os.remove(tmp_pdf_path)

if __name__ == "__main__":
    main()
